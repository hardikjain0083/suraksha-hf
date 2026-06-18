from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from database import db
from datetime import datetime
import uuid
import io
import sys
sys.path.insert(0, ".")
from services.watcher import process_circular, extract_pdf_robust, get_embedder, generate_embeddings
from services.audit_logger import append_audit_log

router = APIRouter()

def get_db():
    return db.client.suraksha_maps

@router.get("/")
async def get_policies():
    database = get_db()
    cursor = database.policies.find({}).sort("valid_from", -1)
    policies = []
    async for p in cursor:
        p["id"] = p.get("policy_id", str(p["_id"]))
        p["_id"] = str(p["_id"])
        if "valid_from" in p and p["valid_from"]:
            p["valid_from"] = p["valid_from"].isoformat() if hasattr(p["valid_from"], 'isoformat') else str(p["valid_from"])
        if "valid_until" in p and p["valid_until"]:
            p["valid_until"] = p["valid_until"].isoformat() if hasattr(p["valid_until"], 'isoformat') else str(p["valid_until"])
        policies.append(p)
    return policies

@router.post("/upload")
async def upload_policy(
    file: UploadFile = File(...),
    title: str = Form(...),
    department: str = Form(...),
    version: str = Form(...)
):
    database = get_db()
    file_bytes = await file.read()
    
    # Extract text based on file type
    ext = file.filename.rsplit(".", 1)[-1].lower()
    text = ""
    
    if ext == "pdf":
        res = await extract_pdf_robust(file_bytes)
        text = res.text
    elif ext in ("docx",):
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to extract text from DOCX: {str(e)}")
    elif ext in ("doc",):
        # Try antiword or fallback to raw text extraction for legacy .doc
        try:
            import subprocess
            result = subprocess.run(
                ["antiword", "-"],
                input=file_bytes, capture_output=True, timeout=10
            )
            if result.returncode == 0:
                text = result.stdout.decode("utf-8", errors="ignore")
            else:
                # Fallback: try treating as docx
                from docx import Document
                doc = Document(io.BytesIO(file_bytes))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception:
            # Last resort: decode bytes
            text = file_bytes.decode("utf-8", errors="ignore")
    else:
        text = file_bytes.decode("utf-8", errors="ignore")
        
    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from document")

    embs = await generate_embeddings([text])
    embedding = embs[0] if embs else []
        
    policy_id = f"POL-{str(uuid.uuid4())[:8].upper()}"
    
    doc = {
        "policy_id": policy_id,
        "title": title,
        "version": version,
        "department_owner_id": department,
        "department": department,
        "status": "active",
        "valid_from": datetime.utcnow(),
        "valid_until": None,
        "content": text,
        "full_text": text,
        "embedding": embedding,
        "covers_clauses": []
    }
    
    await database.policies.insert_one(doc)
    
    await append_audit_log(
        database,
        action_type="POLICY_UPLOADED",
        target_type="policy",
        target_id=policy_id,
        user_id="admin",
        user_name="Administrator",
        details={"title": title, "department": department, "version": version}
    )
    
    return {"status": "success", "policy_id": policy_id}


@router.patch('/{policy_id}/archive')
async def archive_policy(policy_id: str):
    database = get_db()
    res = await database.policies.update_one({'policy_id': policy_id}, {'$set': {'status': 'archived', 'valid_until': datetime.utcnow()}})
    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail='Policy not found')
        
    await append_audit_log(
        database,
        action_type="POLICY_ARCHIVED",
        target_type="policy",
        target_id=policy_id,
        user_id="admin",
        user_name="Administrator"
    )
    
    return {'status': 'success', 'message': 'Policy archived'}

@router.patch('/{policy_id}/unarchive')
async def unarchive_policy(policy_id: str):
    database = get_db()
    res = await database.policies.update_one(
        {'policy_id': policy_id},
        {'$set': {'status': 'active', 'valid_until': None}}
    )
    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail='Policy not found')
        
    await append_audit_log(
        database,
        action_type="POLICY_UNARCHIVED",
        target_type="policy",
        target_id=policy_id,
        user_id="admin",
        user_name="Administrator"
    )
    
    return {'status': 'success', 'message': 'Policy unarchived'}

@router.delete('/{policy_id}')
async def delete_policy(policy_id: str):
    database = get_db()
    res = await database.policies.delete_one({'policy_id': policy_id})
    if res.deleted_count == 0:
        raise HTTPException(status_code=404, detail='Policy not found')
        
    await append_audit_log(
        database,
        action_type="POLICY_DELETED",
        target_type="policy",
        target_id=policy_id,
        user_id="admin",
        user_name="Administrator"
    )
    
    return {'status': 'success', 'message': 'Policy deleted'}

@router.patch('/{policy_id}')
async def update_policy(policy_id: str, title: str = Form(None), version: str = Form(None), department: str = Form(None)):
    database = get_db()
    update_data = {}
    if title: update_data['title'] = title
    if version: update_data['version'] = version
    if department: update_data['department'] = department
    if not update_data:
        raise HTTPException(status_code=400, detail='No fields to update')
        
    res = await database.policies.update_one({'policy_id': policy_id}, {'$set': update_data})
    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail='Policy not found')
        
    await append_audit_log(
        database,
        action_type="POLICY_UPDATED",
        target_type="policy",
        target_id=policy_id,
        user_id="admin",
        user_name="Administrator",
        details=update_data
    )
    
    return {'status': 'success', 'message': 'Policy updated'}
